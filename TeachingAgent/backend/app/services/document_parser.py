from __future__ import annotations

from pathlib import Path
import xml.etree.ElementTree as ET

from docx import Document
import fitz  # PyMuPDF
from sqlalchemy.orm import Session

from app.models.db import ParsedDocument, SourceFile


class DocumentParserService:
    def __init__(self, db: Session):
        self.db = db

    def parse_file(self, source_file: SourceFile, absolute_path: Path) -> str:
        extension = source_file.extension.lower()
        if extension == ".pdf":
            return self._parse_pdf(absolute_path)
        if extension == ".docx":
            return self._parse_docx(absolute_path)
        if extension in {".xml", ".bpml"}:
            return self._parse_xml(absolute_path)
        if extension in {".ddf", ".mxl", ".txo"}:
            return self._parse_text(absolute_path)
        if extension in {".png", ".jpg", ".jpeg"}:
            return ""
        return self._parse_text(absolute_path)

    def persist(self, source_file: SourceFile, full_text: str) -> None:
        record = ParsedDocument(source_file_id=source_file.id, full_text=full_text)
        self.db.add(record)
        self.db.commit()

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        reader = PdfReader(str(path))
        pages = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append(f"[Page {page_index}]\n{text}")
        return "\n\n".join(pages)

    @staticmethod
    def _parse_docx(path: Path) -> str:
        doc = Document(str(path))
        lines = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return "\n".join(lines)

    @staticmethod
    def _parse_xml(path: Path) -> str:
        tree = ET.parse(str(path))
        root = tree.getroot()
        text_bits: list[str] = []
        for element in root.iter():
            if element.text and element.text.strip():
                text_bits.append(element.text.strip())
            for key, value in element.attrib.items():
                text_bits.append(f"{key}: {value}")
        return "\n".join(text_bits)

    @staticmethod
    def _parse_text(path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1", errors="ignore")
